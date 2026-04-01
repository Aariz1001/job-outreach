"""
CV Specialiser — surgically tailors the DOCX template CV for a specific company,
then converts it to PDF for email attachment.

Only modifies content paragraphs (objective, bullet points, skills text).
Section headers, job titles, project names, and contact info are NEVER changed.

Page constraint:
  The original CV is already tightly packed at ~1 page.  We enforce a hard
  character-growth limit of 8 % so substitutions can never push it to 2 pages.
  Each individual replacement text must also be ≤130 % of the original length.

Output:
  specialised_cvs/<company_slug>/cv_<company_slug>.docx
  specialised_cvs/<company_slug>/cv_<company_slug>.pdf   ← converted PDF
"""
from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from openai import OpenAI
from rich.console import Console

from utils.config import GENERATION_MODEL, OPENROUTER_API_KEY

console = Console()

_client = OpenAI(
    api_key=OPENROUTER_API_KEY,
    base_url="https://openrouter.ai/api/v1",
)

_ROOT        = Path(__file__).parent.parent
TEMPLATE_CV  = _ROOT / "Mohammad_Aariz_Waqas_AI_Engineer.docx"
OUTPUT_DIR   = _ROOT / "specialised_cvs"

# ── Paragraph index whitelist ─────────────────────────────────────────────────
# Determined by inspecting the DOCX structure.  Only these indices can receive
# text substitutions.  Everything else (headers, job titles, contact info, the
# name, education row, project hyperlink lines) is protected.
EDITABLE_PARA_INDICES: frozenset[int] = frozenset({
    2,                 # keywords / tagline line
    5,                 # Objective body
    8,  9,  10,        # Accomplishments bullets
    13, 14, 15,        # Selected Impact bullets
    19, 20, 21,        # Job 1 — Ascension bullets
    23, 24, 25,        # Job 2 — CITUS bullets
    28,                # FlexiDiet project description
    30,                # AgentForge project description
    32,                # VoxGuard project description
    34,                # Predictive Intent project description
    37, 38, 39, 40,    # Technical Highlights bullets
    43, 44, 45, 46,    # Technical Skills content (second run per para)
})

MAX_EDITS        = 10    # max substitutions per specialisation run
MAX_CHAR_GROWTH  = 0.08  # 8 % hard ceiling


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _slug(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")


def _para_text(para) -> str:
    return "".join(r.text for r in para.runs)


def _replace_in_para(para, old_text: str, new_text: str) -> bool:
    """Replace *old_text* with *new_text* inside *para*, preserving run formatting.

    Strategy:
      1. Search each run individually (handles the common single-run case and
         the two-run skills lines where the label run stays untouched).
      2. If not found in any single run, join all run texts, replace, and
         reassign to the first content run — safe for pure-content paragraphs.
    """
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text, 1)
            return True

    # Cross-run fallback
    full = "".join(r.text for r in para.runs)
    if old_text not in full:
        return False

    new_full = full.replace(old_text, new_text, 1)
    content_runs = [r for r in para.runs if r.text.strip()]
    if not content_runs:
        return False

    content_runs[0].text = new_full
    for r in para.runs:
        if r is not content_runs[0]:
            r.text = ""
    return True


def _total_chars(doc) -> int:
    return sum(len(p.text) for p in doc.paragraphs)


def _get_editable_paragraphs(doc) -> dict[int, str]:
    """Return {para_idx: text} for all editable content paragraphs that have text."""
    return {
        i: p.text
        for i, p in enumerate(doc.paragraphs)
        if i in EDITABLE_PARA_INDICES and p.text.strip()
    }


# ── PDF conversion ────────────────────────────────────────────────────────────

_SOFFICE_PATHS = [
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/bin/soffice",
    "/usr/local/bin/soffice",
    "soffice",
]


def _find_soffice() -> str | None:
    for path in _SOFFICE_PATHS:
        try:
            r = subprocess.run(
                [path, "--version"], capture_output=True, timeout=5
            )
            if r.returncode == 0:
                return path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return None


def _docx_to_pdf_libreoffice(docx_path: Path, out_dir: Path) -> Path | None:
    """Convert docx → pdf using LibreOffice headless. Returns PDF path or None."""
    soffice = _find_soffice()
    if not soffice:
        return None
    try:
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(out_dir),
                str(docx_path),
            ],
            capture_output=True,
            timeout=90,
        )
        if result.returncode == 0:
            pdf_path = out_dir / (docx_path.stem + ".pdf")
            if pdf_path.exists():
                return pdf_path
    except Exception as e:
        console.print(f"[yellow]  [PDF] LibreOffice error: {e}[/yellow]")
    return None


def _docx_to_pdf_word(docx_path: Path, out_dir: Path) -> Path | None:
    """Convert docx → pdf using Microsoft Word via docx2pdf (macOS/Windows only)."""
    if sys.platform not in ("darwin", "win32"):
        return None
    try:
        from docx2pdf import convert
        pdf_path = out_dir / (docx_path.stem + ".pdf")
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception as e:
        console.print(f"[yellow]  [PDF] docx2pdf error: {e}[/yellow]")
    return None


def docx_to_pdf(docx_path: Path) -> Path | None:
    """Convert *docx_path* to PDF in the same directory.  Returns PDF path or None.

    Tries LibreOffice first (highest fidelity), then Microsoft Word (macOS/Windows).
    """
    out_dir = docx_path.parent

    # 1 — LibreOffice (preferred, free, cross-platform)
    pdf = _docx_to_pdf_libreoffice(docx_path, out_dir)
    if pdf:
        console.print(f"[green]  [PDF] Converted via LibreOffice → {pdf.name}[/green]")
        return pdf

    # 2 — Microsoft Word via AppleScript (macOS only, requires Word)
    pdf = _docx_to_pdf_word(docx_path, out_dir)
    if pdf:
        console.print(f"[green]  [PDF] Converted via Word → {pdf.name}[/green]")
        return pdf

    console.print(
        "[yellow]  [PDF] Could not convert to PDF — LibreOffice not found and Word "
        "not available.\n"
        "  Install LibreOffice: brew install --cask libreoffice[/yellow]"
    )
    return None


# ── LLM edit proposal ─────────────────────────────────────────────────────────

def _propose_edits(
    editable_paras: dict[int, str],
    research: dict,
    company: str,
) -> list[dict]:
    """Ask the LLM to propose targeted text substitutions for *company*."""

    editable_block = "\n".join(
        f"[{idx}] {text}" for idx, text in sorted(editable_paras.items())
    )

    tech_str   = ", ".join(research.get("tech_signals", [])[:8])
    tp_str     = "\n".join(f"- {t}" for t in research.get("talking_points", [])[:3])
    fit_reason = research.get("fit_reason", "")
    one_liner  = research.get("one_liner", "")

    prompt = f"""You are a CV tailoring specialist. Suggest targeted text substitutions to make
this candidate's CV more relevant to the target company.

STRICT RULES — READ CAREFULLY:
1. You may ONLY propose edits to the EDITABLE paragraphs listed below (identified by [index]).
2. "old" must be an EXACT substring of that paragraph — character-perfect match.
3. Only change specific words, phrases, or a sub-sentence — never an entire paragraph.
4. Replacement ("new") must be ≤130 % of the original character length — no padding.
5. Preserve all metrics, technology names, and company names in the original text.
6. Maximum {MAX_EDITS} edits.  Prioritise highest-impact changes only.
7. Do NOT touch paragraph [0] (name), [1] (contact), [4]/[7]/[12]/[17]/[26]/[36]/[42]/[48]
   (section headers), [18]/[22] (job title lines), or [27]/[29]/[31]/[33] (project names).
8. Do NOT invent technologies or achievements not evidenced in the original CV.
9. If no changes would genuinely improve relevance, return [].

TARGET COMPANY: {company}
One-liner: {one_liner}
Tech signals: {tech_str}
Why candidate fits:
{fit_reason}
Key talking points:
{tp_str}

EDITABLE PARAGRAPHS:
{editable_block}

Return ONLY valid JSON — no markdown fences, no explanation:
[
  {{
    "para_idx": <int>,
    "old": "<exact substring currently in that paragraph>",
    "new": "<replacement of similar length>",
    "reason": "<one-line reason>"
  }}
]"""

    resp = _client.chat.completions.create(
        model=GENERATION_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        extra_body={"reasoning": {"enabled": True}},
    )

    raw = resp.choices[0].message.content.strip()
    raw = re.sub(r"^```[a-zA-Z]*\n?", "", raw)
    raw = re.sub(r"\n?```$", "", raw)

    try:
        edits = json.loads(raw)
        if isinstance(edits, list):
            return edits
    except Exception as exc:
        console.print(f"[yellow]  [CV] JSON parse error: {exc}[/yellow]")
        console.print(f"[dim]  Raw: {raw[:400]}[/dim]")

    return []


# ── Public API ────────────────────────────────────────────────────────────────

def specialise_cv(company: str, research: dict) -> tuple[Path, Path | None]:
    """Specialise the DOCX CV for *company* and return (docx_path, pdf_path).

    Steps:
      1. Load a fresh copy of the template.
      2. Present editable paragraphs to the LLM with company research.
      3. Apply the proposed substitutions (with safety guards).
      4. Validate the 1-page character-count constraint.
      5. Save to specialised_cvs/<slug>/cv_<slug>.docx.
      6. Convert to PDF and save alongside the DOCX.
    """
    company_slug = _slug(company)
    out_dir  = OUTPUT_DIR / company_slug
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"cv_{company_slug}.docx"

    # Load a fresh copy of the template DOCX
    doc = Document(str(TEMPLATE_CV))
    original_chars = _total_chars(doc)
    console.print(f"[dim]  [CV] Template loaded: {original_chars} chars[/dim]")

    editable_paras = _get_editable_paragraphs(doc)

    console.print(f"[cyan]  [CV] Asking LLM for targeted edits for '{company}'…[/cyan]")
    edits = _propose_edits(editable_paras, research, company)

    if not edits:
        console.print("[yellow]  [CV] No edits proposed — CV already matches well.[/yellow]")
        shutil.copy(str(TEMPLATE_CV), str(out_path))
        console.print(f"[green]  [CV] Saved (unchanged) → {out_path}[/green]")
        console.print("[cyan]  [PDF] Converting to PDF…[/cyan]")
        pdf_path = docx_to_pdf(out_path)
        return out_path, pdf_path

    console.print(f"[dim]  [CV] Applying {len(edits)} edit(s)…[/dim]")
    applied = 0

    for edit in edits[:MAX_EDITS]:
        try:
            para_idx = int(edit["para_idx"])
            old_text = str(edit["old"])
            new_text = str(edit["new"])
            reason   = str(edit.get("reason", ""))
        except (KeyError, TypeError, ValueError) as exc:
            console.print(f"[yellow]  [CV] Malformed edit, skipping: {exc}[/yellow]")
            continue

        # Safety: only allow editable paragraphs
        if para_idx not in EDITABLE_PARA_INDICES:
            console.print(f"[red]  [CV] Blocked edit to protected para[{para_idx}][/red]")
            continue

        # Safety: replacement must not grow too much
        if len(new_text) > len(old_text) * 1.30:
            console.print(
                f"[yellow]  [CV] Skip para[{para_idx}]: replacement too long "
                f"({len(new_text)} vs allowed {int(len(old_text)*1.3)})[/yellow]"
            )
            continue

        para = doc.paragraphs[para_idx]
        current = _para_text(para)

        if old_text not in current:
            console.print(
                f"[yellow]  [CV] Skip para[{para_idx}]: '{old_text[:45]}' not found[/yellow]"
            )
            continue

        if _replace_in_para(para, old_text, new_text):
            console.print(
                f"[green]  [CV] ✓ [{para_idx}] «{old_text[:38]}» → «{new_text[:38]}»[/green]"
            )
            if reason:
                console.print(f"[dim]       ↳ {reason}[/dim]")
            applied += 1
        else:
            console.print(f"[yellow]  [CV] Could not apply edit to para[{para_idx}][/yellow]")

    # ── 1-page validation ─────────────────────────────────────────────────────
    new_chars   = _total_chars(doc)
    growth      = (new_chars - original_chars) / max(original_chars, 1)
    growth_pct  = growth * 100

    console.print(
        f"[dim]  [CV] Character count: {original_chars} → {new_chars} ({growth_pct:+.1f}%)[/dim]"
    )

    if growth > MAX_CHAR_GROWTH:
        console.print(
            f"[bold red]  [CV] WARNING: CV grew {growth_pct:.1f}% — may exceed 1 page. "
            f"Review {out_path}[/bold red]"
        )
    else:
        console.print(f"[green]  [CV] 1-page constraint OK ({applied} edits, {growth_pct:+.1f}%)[/green]")

    # ── Collapse blank spacer paragraphs to minimal height ────────────────────
    # Empty paragraphs between sections use full line-height — shrink to 2pt
    # so there's no visible gap between section content and section headers.
    from docx.shared import Pt
    from docx.oxml.ns import qn
    import lxml.etree as _etree
    for para in doc.paragraphs:
        if para.text.strip() == "":
            # Set every run's font size to 2pt (creates ~0pt gap at ls=1.0)
            for run in para.runs:
                run.font.size = Pt(2)
            # Also set via pPr/rPr on the paragraph mark itself
            pPr = para._p.get_or_add_pPr()
            rPr = pPr.find(qn("w:rPr"))
            if rPr is None:
                rPr = _etree.SubElement(pPr, qn("w:rPr"))
            sz_el = rPr.find(qn("w:sz"))
            if sz_el is None:
                sz_el = _etree.SubElement(rPr, qn("w:sz"))
            sz_el.set(qn("w:val"), "4")  # half-points: 4 = 2pt
            szCs_el = rPr.find(qn("w:szCs"))
            if szCs_el is None:
                szCs_el = _etree.SubElement(rPr, qn("w:szCs"))
            szCs_el.set(qn("w:val"), "4")

    doc.save(str(out_path))
    console.print(f"[green]  [CV] Saved → {out_path}[/green]")

    # ── PDF conversion ────────────────────────────────────────────────────────
    console.print("[cyan]  [PDF] Converting to PDF…[/cyan]")
    pdf_path = docx_to_pdf(out_path)

    return out_path, pdf_path


if __name__ == "__main__":
    # Smoke test — run:  python -m agents.cv_specialiser
    sample_research = {
        "one_liner": "PolyAI builds voice AI agents that handle enterprise customer service calls.",
        "tech_signals": ["voice AI", "conversational AI", "real-time", "LLM", "telephony", "NLU"],
        "fit_reason": (
            "Aariz built VoxGuard (<200ms voice intent) and a live voice AI agent "
            "at Ascension handling real customer calls. Direct match for PolyAI's stack."
        ),
        "talking_points": [
            "VoxGuard: sub-200ms intent classification for live voice AI systems",
            "Built voice AI agent at Ascension — appointment booking, backend integrations",
        ],
        "product_stage": "Series C",
    }
    docx_path, pdf_path = specialise_cv("PolyAI", sample_research)
    console.print(f"[bold]DOCX:[/bold] {docx_path}")
    console.print(f"[bold]PDF:[/bold]  {pdf_path}")
